from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ── Global Style ──
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.5

# ── Helper Functions ──
def add_page_break():
    doc.add_page_break()

def set_margins(top=1, bottom=1, left=1.25, right=1.25):
    for section in doc.sections:
        section.top_margin = Inches(top)
        section.bottom_margin = Inches(bottom)
        section.left_margin = Inches(left)
        section.right_margin = Inches(right)

set_margins()

def add_title_page():
    for _ in range(6):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('PROJECT SYNOPSIS')
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = RGBColor(0, 0, 0)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('ElevoraAI')
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0, 51, 102)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('An AI-Powered Smart Education and Career Advancement Platform')
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0, 51, 102)
    
    for _ in range(4):
        doc.add_paragraph()
    
    info_items = [
        ('Project Type', 'Full-Stack Web Application'),
        ('Frontend', 'React 18, TypeScript, Tailwind CSS, Vite'),
        ('Backend', 'Node.js, Express 5, MongoDB, Mongoose'),
        ('Authentication', 'JWT, Firebase Google OAuth, OTP Email Verification'),
        ('AI Integration', 'OpenAI (GPT-4), Google Gemini, Groq (LLaMA 3)'),
        ('Real-Time', 'Socket.IO'),
        ('Payments', 'Stripe'),
        ('Code Compiler', 'Judge0 API'),
    ]
    for label, value in info_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f'{label}: ')
        run.bold = True
        run.font.size = Pt(12)
        run = p.add_run(value)
        run.font.size = Pt(12)
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Open Source Project under Code Social')
    run.italic = True
    run.font.size = Pt(12)
    
    add_page_break()

def add_heading_styled(text, level=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 0 else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(18 if level <= 1 else 14)
    run.font.color.rgb = RGBColor(0, 51, 102)
    if level <= 1:
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(12)
    return p

def add_subheading(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0, 51, 102)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(6)
    return p

def add_body(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.size = Pt(12)
    return p

def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(12)
        run = p.add_run(text)
        run.font.size = Pt(12)
    else:
        run = p.add_run(text)
        run.font.size = Pt(12)
    return p

# ══════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════
add_title_page()

# ══════════════════════════════════════════════════
# TABLE OF CONTENTS
# ══════════════════════════════════════════════════
add_heading_styled('Table of Contents', 0)
add_page_break()

# ══════════════════════════════════════════════════
# 1. INTRODUCTION
# ══════════════════════════════════════════════════
add_heading_styled('1. Introduction', 1)

add_body(
    'ElevoraAI is a comprehensive, full-stack, AI-powered smart education and career advancement platform '
    'designed to empower students, developers, and job seekers in their technical learning and career growth journey. '
    'The platform serves as a centralized ecosystem that integrates structured learning paths, AI-driven personal assistance, '
    'coding practice tools, placement preparation resources, resume and cover letter building utilities, community collaboration '
    'features, and a powerful administrative control panel — all accessible through a unified, role-based interface.'
)

add_body(
    'The platform addresses the fundamental challenge of fragmented learning and career preparation tools. In the current landscape, '
    'a student must navigate multiple platforms for different aspects of skill development. Online courses are taken on platforms like '
    'Coursera or Udemy. Coding practice happens on LeetCode, HackerRank, or Codeforces. Resume building requires separate tools like '
    'Canva or Novoresume. Job searching is conducted through LinkedIn, Naukri, or Internshala. Community interaction occurs on Stack Overflow '
    'or Discord. This fragmented approach results in significant inefficiency, loss of contextual progress tracking, data silos, and a '
    'disjointed user experience that hinders effective learning and career preparation.'
)

add_body(
    'ElevoraAI eliminates this fragmentation by bringing all these capabilities into a single, cohesive platform. '
    'Built using a modern technology stack, the frontend leverages React 18 with TypeScript for type-safe, component-based UI development, '
    'Vite for fast builds and hot module replacement, and Tailwind CSS with Radix UI primitives for a polished, responsive interface. '
    'The backend is powered by Node.js with Express 5, using MongoDB with Mongoose for flexible document-based data storage. '
    'Authentication is secured through JWT tokens stored in httpOnly cookies, with support for email OTP verification, password reset flows, '
    'and Firebase Google OAuth integration. The platform also integrates multiple AI services: Google Gemini powers the AI Study Buddy chatbot, '
    'quiz generation, and note generation; OpenAI GPT-4 handles skill assessment generation and evaluation; and Groq with LLaMA 3 drives '
    'the project recommendation engine. Real-time capabilities are implemented using Socket.IO for live contest leaderboards and coding arena features. '
    'Code execution is handled through the Judge0 API, which supports compilation and testing across multiple programming languages. '
    'Payment processing is managed through Stripe integration for premium subscription plans.'
)

add_body(
    'The platform includes over twenty major feature modules. The authentication module supports role-based access with separate user and admin '
    'interfaces. The smart user dashboard displays personalized analytics including streak data, learning progress, achievements, and quick-access '
    'shortcuts. The learning hub provides four structured tracks covering Data Structures and Algorithms, Java, MERN Stack, and AI/Machine Learning, '
    'each with module-based content organization and progress tracking. The AI Study Buddy chatbot offers 24x7 doubt resolution and career guidance. '
    'The quiz system supports both manually created and AI-generated quizzes. Skill assessments span five domains with AI-generated questions and '
    'detailed feedback. The coding platform features the Monaco Editor with multi-language compilation support. Contest management includes '
    'real-time leaderboards and Elo-based rating calculations. The resume builder includes ATS scanning and PDF export. Additional modules include '
    'a tech news feed, community Q&A forum, notes system, video course platform, placement preparation hub, career roadmaps, calendar and task '
    'management, multi-mode calculator, premium subscriptions, admin panel with analytics, notification system, and global search.'
)

add_body('The key features of the platform are summarized below:')
add_bullet('Login and registration with email OTP verification and Google OAuth authentication.', 'Secure Authentication: ')
add_bullet('Separate dashboards and access controls for regular users and platform administrators.', 'Role-Based Access: ')
add_bullet('Structured learning tracks for DSA, Java, MERN Stack, and AI/ML with module-level progress tracking.', 'Learning Hub: ')
add_bullet('24x7 AI-powered chatbot for doubt resolution, career advice, and learning assistance using Google Gemini.', 'AI Study Buddy: ')
add_bullet('Monaco Editor-based coding platform with Judge0 compiler integration and multiple language support.', 'Coding Platform: ')
add_bullet('Real-time coding contests with Socket.IO-powered leaderboards and Elo rating system.', 'Contest Management: ')
add_bullet('AI-generated and manually created quizzes with timed attempts, scoring, and performance history.', 'Quiz System: ')
add_bullet('Skill assessments across five tracks with AI-generated questions and personalized feedback.', 'Skill Assessments: ')
add_bullet('ATS-compliant resume builder with PDF export and cover letter generator.', 'Resume and Cover Letter Builder: ')
add_bullet('Tech news aggregation, job listings, internship updates, and career resources.', 'Tech Feed and Career Updates: ')
add_bullet('Community Q&A forum for peer-to-peer learning with admin moderation.', 'Community Forum: ')
add_bullet('AI-powered project recommendations based on user preferences and skill levels.', 'Project Recommender: ')
add_bullet('Comprehensive notes system with AI-generated notes and structured subject-wise content.', 'Notes System: ')
add_bullet('YouTube-based video course platform with progress tracking and bookmarking.', 'Video Course Platform: ')
add_bullet('Placement preparation module with job opportunities, interview resources, and DSA practice.', 'Placement Prep: ')
add_bullet('Career roadmaps, mock interview preparation, and interview experience sharing.', 'Interview and Roadmaps: ')
add_bullet('Calendar view, task management, multi-mode calculator, and global search.', 'Productivity Tools: ')
add_bullet('Stripe-integrated premium subscription plans with feature comparison.', 'Premium and Payments: ')
add_bullet('Admin dashboard for user management, content moderation, analytics, and system settings.', 'Admin Panel: ')
add_bullet('In-app notifications and email notifications for platform updates and milestones.', 'Notifications: ')

add_page_break()

# ══════════════════════════════════════════════════
# 2. PROBLEM STATEMENT
# ══════════════════════════════════════════════════
add_heading_styled('2. Problem Statement', 1)

add_body(
    'Students and early-career developers face a significant and persistent challenge in navigating the modern technology landscape. '
    'The process of acquiring technical skills, preparing for job placements, and building a professional profile is highly fragmented '
    'across multiple disconnected platforms and tools. This fragmentation creates substantial inefficiencies and barriers to effective '
    'learning and career advancement.'
)

add_body(
    'The educational technology ecosystem currently operates in silos. Online learning platforms such as Coursera, Udemy, and edX provide '
    'comprehensive course libraries but lack integrated coding environments where learners can practice concepts in real-time. Coding practice '
    'platforms like LeetCode, HackerRank, and Codeforces offer excellent problem-solving environments but do not provide structured curricula '
    'or AI-powered learning assistance. Resume building tools like Canva, Novoresume, and Zety focus solely on document design without '
    'integrating skill tracking or placement preparation. Job search platforms such as LinkedIn, Naukri, and Internshala connect candidates '
    'with employers but offer no support for skill development or interview preparation. Community platforms like Stack Overflow and Discord '
    'facilitate knowledge sharing but lack integration with formal learning resources or progress tracking.'
)

add_body(
    'This fragmented landscape creates several critical problems for learners. First, users must maintain accounts and profiles across '
    'multiple platforms, leading to data duplication and inconsistent progress tracking. Second, context switching between platforms '
    'consumes significant time and mental energy, reducing the efficiency of study sessions. Third, there is no unified mechanism to '
    'track overall skill development — a user may complete courses on one platform, solve problems on another, and build projects on a third, '
    'with no single view of their complete learning journey. Fourth, existing platforms rarely provide AI-powered personalized guidance '
    'that adapts to individual learning styles, pace, and career goals. Fifth, placement preparation—a critical need for students—requires '
    'coordinated effort across resume building, skill assessment, aptitude practice, technical interview preparation, and job search, '
    'which currently cannot be managed from a single interface.'
)

add_body(
    'From an administrative perspective, educational institutions and program coordinators lack unified tools to monitor learner progress, '
    'manage content, analyze platform usage, and moderate community interactions. Without an integrated admin panel, managing a large '
    'learner base requires manual effort and fragmented reporting tools.'
)

add_body(
    'Furthermore, existing platforms lack several important features that modern learners need. AI-powered tools for automated content '
    'generation, personalized recommendations, and intelligent tutoring are still absent from most platforms. Real-time collaboration '
    'features such as live coding contests, shared leaderboards, and community problem-solving are typically limited to dedicated '
    'competition platforms. Role-based access control that differentiates between learners and administrators is not a standard feature '
    'in most educational tools. Additionally, few platforms offer comprehensive career support including ATS-optimized resume building, '
    'cover letter generation, mock interview practice, and placement preparation in an integrated manner.'
)

add_body(
    'The core problems addressed by ElevoraAI are summarized as follows:'
)
add_bullet('Learners must use separate platforms for courses, coding practice, resume building, job search, and community interaction.', 'Platform Fragmentation: ')
add_bullet('No unified view of overall skill development and learning progress across different activities.', 'Disjointed Progress Tracking: ')
add_bullet('Most platforms lack AI-powered personalization, adaptive learning paths, and intelligent tutoring capabilities.', 'Lack of AI Integration: ')
add_bullet('Existing tools do not combine real-time coding contests, leaderboards, and collaborative features with structured learning.', 'Limited Real-Time Collaboration: ')
add_bullet('Administrators lack unified dashboards for user management, content moderation, analytics, and platform oversight.', 'No Admin Control: ')
add_bullet('Career preparation requires coordinated use of resume builders, skill assessments, interview practice, and job search tools.', 'Career Preparation Disconnect: ')
add_bullet('Few platforms offer role-based access control that securely differentiates between learner and administrator privileges.', 'Access Control Gaps: ')
add_bullet('Automated content generation, AI-assisted doubt resolution, and personalized recommendations are not widely available.', 'Absence of AI-Assisted Tools: ')

add_page_break()

# ══════════════════════════════════════════════════
# 3. LITERATURE SURVEY
# ══════════════════════════════════════════════════
add_heading_styled('3. Literature Survey', 1)

add_body(
    '1. Limitations of Current E-Learning Platforms — A comprehensive review of existing e-learning platforms including Coursera, Udemy, '
    'and edX reveals that while these platforms offer extensive course libraries and structured content delivery, they lack several critical '
    'features. Most notably, they do not provide integrated coding environments for hands-on practice, AI-powered doubt resolution for '
    'real-time assistance, or placement-specific preparation modules. Users must supplement their learning with additional platforms for '
    'coding practice, skill assessment, and career preparation, resulting in a fragmented and inefficient learning experience. The absence '
    'of role-based access control in most platforms also limits their utility for institutional deployment where administrator oversight '
    'is required.'
)

add_body(
    '2. Artificial Intelligence in Educational Technology — Research by Mollick and Mollick (2023) and others demonstrates that '
    'AI-powered chatbots and personalized recommendation systems significantly improve learner engagement, knowledge retention, and '
    'learning outcomes. The integration of large language models such as GPT-4, Gemini, and LLaMA for real-time doubt resolution, '
    'automated content generation, and intelligent assessment represents a growing trend in intelligent tutoring systems. Studies indicate '
    'that students who use AI-powered learning assistants show improved problem-solving skills and higher motivation compared to those '
    'using traditional learning methods alone. The use of multiple AI models for different purposes — conversational assistance, content '
    'generation, and skill evaluation — represents a best practice in educational AI architecture.'
)

add_body(
    '3. Role-Based Access Control in Web Applications — Research on role-based access control (RBAC) implementation in educational '
    'platforms demonstrates that separating user and administrator privileges is essential for system security, data integrity, and '
    'efficient content management at scale. The implementation of JSON Web Token (JWT) based authentication with role middleware '
    'is widely adopted as an industry best practice. Studies by Sandhu et al. (1996) on RBAC models provide the foundational framework '
    'that modern educational platforms use to manage permissions, protect sensitive data, and enable granular control over platform features.'
)

add_body(
    '4. Gamification and Real-Time Collaboration in Learning — Literature on gamified learning environments by Hamari, Koivisto, and '
    'Sarsa (2014) indicates that gamification elements such as leaderboards, achievement badges, streak tracking, and competitive contests '
    'significantly increase user motivation, engagement, and knowledge retention. Real-time technologies like WebSockets and Socket.IO '
    'enable live collaboration features that enhance the social learning experience. Studies show that competitive elements combined with '
    'collaborative features create an optimal environment for skill development, particularly in technical domains like programming and '
    'data structures.'
)

add_body(
    '5. Applicant Tracking Systems and Resume Optimization — Research on applicant tracking systems (ATS) indicates that over 75 percent '
    'of resumes are rejected before reaching a human recruiter due to poor keyword optimization, formatting issues, and missing sections. '
    'Automated resume scanning tools that analyze keyword density, section completeness, role-specific terminology, and ATS compliance '
    'significantly improve interview call rates. Studies recommend that resume builders incorporate real-time ATS scoring and suggestions '
    'to help job seekers optimize their applications for specific roles and industries.'
)

add_body(
    '6. Open Source Contribution as a Learning Model — Research in software engineering education highlights that contributing to '
    'open-source projects provides invaluable real-world experience in version control, code review, collaborative development, and '
    'project management. Platforms that facilitate open-source contribution alongside structured learning create a more holistic skill '
    'development environment. Studies by Pinto et al. (2016) demonstrate that students who participate in open-source projects show '
    'significant improvement in code quality, technical communication, and confidence in software development practices.'
)

add_page_break()

# ══════════════════════════════════════════════════
# 4. OBJECTIVES
# ══════════════════════════════════════════════════
add_heading_styled('4. Objectives', 1)

add_body('The primary objectives of the ElevoraAI platform are as follows:\n')

add_bullet('To develop a unified, AI-powered platform that integrates learning, skill assessment, coding practice, and career preparation in a single system.')
add_bullet('To provide structured learning tracks in Data Structures and Algorithms, Java, MERN Stack, and AI/Machine Learning with module-based content and progress tracking.')
add_bullet('To implement an intelligent AI Study Buddy chatbot using Google Gemini for 24x7 doubt resolution, resource suggestions, and career guidance.')
add_bullet('To build a fully functional coding platform with Monaco Editor integration, multi-language compilation via Judge0 API, and contest management with real-time leaderboards.')
add_bullet('To create an ATS-compliant resume builder and cover letter generator with PDF export and keyword optimization features.')
add_bullet('To develop an AI-driven skill assessment engine that generates personalized tests across five domains and provides detailed feedback with improvement recommendations.')
add_bullet('To integrate a tech news feed, job listings, internship opportunities, and career resources using external APIs and curated content.')
add_bullet('To build a community Q&A forum for peer-to-peer knowledge sharing with admin moderation and content management capabilities.')
add_bullet('To implement a comprehensive notes system with AI-generated notes, structured subject content, and user-created notes with search and categorization.')
add_bullet('To develop a video course platform with YouTube API integration, progress tracking, bookmarking, and continue-learning features.')
add_bullet('To implement a full-featured admin panel with dashboards for user management, content moderation, quiz management, analytics, system logs, and configuration settings.')
add_bullet('To support secure, role-based authentication with JWT tokens, email OTP verification, Firebase Google OAuth, and password reset flows.')
add_bullet('To integrate real-time Socket.IO communication for live contest rooms, leaderboard updates, and participant tracking.')
add_bullet('To implement a notification system for in-app alerts and email notifications for platform updates, milestones, and reminders.')
add_bullet('To provide productivity tools including calendar view, task management, multi-mode calculator, and global search functionality.')

add_page_break()

# ══════════════════════════════════════════════════
# 5. METHODOLOGY / FLOW OF WORK
# ══════════════════════════════════════════════════
add_heading_styled('5. Methodology / Flow of Work', 1)

add_body(
    'The ElevoraAI platform follows a modular client-server architecture with a RESTful API design pattern. The frontend and backend '
    'are developed as separate, independent services that communicate through well-defined HTTP endpoints. The architecture ensures '
    'separation of concerns, scalability, and maintainability.'
)

add_subheading('5.1 Authentication Flow')
add_body(
    'New users register using their email address, name, and password. The system sends a one-time password (OTP) to the provided '
    'email address for verification. Upon successful OTP verification, the user account is created and a JWT token is issued. '
    'Users may also register and log in using Firebase Google OAuth, which eliminates the need for password management. '
    'The system assigns a role — either user or admin — during registration. All protected routes validate the JWT token through '
    'middleware, and admin-only routes perform an additional role check. Session management uses httpOnly cookies for enhanced security. '
    'Password reset functionality sends a reset link to the registered email address with a time-limited token.'
)

add_subheading('5.2 User Workflow')
add_body(
    'Upon successful authentication, users are directed to a personalized dashboard. The dashboard displays a welcome message, '
    'current streak count, total points, skill level, learning progress summaries, recent achievements, and quick-access shortcuts '
    'to all major modules. The dashboard fetches data from multiple backend endpoints to provide a comprehensive overview of the '
    'user\'s platform activity. From the dashboard, users can navigate to any module using the navigation bar or sidebar.'
)

add_subheading('5.3 Learning Hub Workflow')
add_body(
    'The Learning Hub presents four structured tracks: DSA, Java, MERN Stack, and AI/ML. Each track contains pre-defined modules '
    'with multiple topics. Users can view module content, mark topics as complete, write personal notes, submit ratings and reviews, '
    'and track their overall progress within each track. Progress data is persisted to the backend and synchronized across sessions. '
    'The Learning Hub also provides links to external practice resources such as GeeksforGeeks, LeetCode, and HackerRank.'
)

add_subheading('5.4 AI-Powered Features Workflow')
add_body(
    'The AI Study Buddy chatbot accepts natural language queries across four categories: learning help, career advice, quiz topics, '
    'and general chat. Queries are sent to the backend which forwards them to the Google Gemini API. Responses are formatted with '
    'Markdown rendering and syntax-highlighted code blocks. The AI Quiz Generator allows users to specify a topic, difficulty level, '
    'and question type, and generates a complete quiz using Gemini. Skill Assessments use OpenAI to generate domain-specific questions '
    'across five tracks, evaluate user responses, and return scores, accuracy percentages, skill levels, improvement areas, and '
    'personalized learning roadmaps. The Project Recommender uses the Groq API with LLaMA 3 to suggest projects based on user preferences '
    'including tech stack, career focus, skill level, and interests.'
)

add_subheading('5.5 Coding and Contest Workflow')
add_body(
    'The coding platform integrates the Monaco Editor with support for multiple programming languages. Users can write, test, and '
    'submit code through the Judge0 API for compilation and execution. The contest system supports the full lifecycle of coding '
    'competitions: administrators create contests with problems and time limits, users register for upcoming contests, participants '
    'solve problems within the time limit, submissions are evaluated in real-time, leaderboards are updated dynamically via Socket.IO, '
    'and Elo ratings are calculated based on performance.'
)

add_subheading('5.6 Admin Workflow')
add_body(
    'Admin users access a dedicated dashboard with ten management tabs. The overview tab displays platform statistics including '
    'total users, active courses, and feedback counts. The user management tab allows CRUD operations on registered users. '
    'The content management tab enables adding, editing, and deleting courses with YouTube video integration. The quiz management '
    'tab provides full control over quiz creation and question management. The community tab allows moderation of forum questions '
    'and answers. The analytics tab uses Chart.js to visualize platform usage data. The system logs tab provides an audit trail '
    'of administrative actions. The settings tab allows configuration of system parameters such as site name, maintenance mode, '
    'and registration settings.'
)

add_subheading('5.7 Major Modules')
add_body('The platform is organized into the following major modules:\n')

modules_list = [
    'Authentication Module — Handles user registration, login, OTP verification, Google OAuth, password management, and session handling using JWT tokens.',
    'User Dashboard Module — Aggregates and displays personalized user data including streaks, progress, achievements, and platform activity summaries.',
    'Learning Hub Module — Provides structured learning tracks for DSA, Java, MERN Stack, and AI/ML with module-level progress tracking.',
    'AI Study Buddy Module — Google Gemini-powered chatbot for real-time doubt resolution, career guidance, and learning assistance.',
    'Quiz Module — Supports manual quiz creation by administrators and AI-generated quizzes with timed attempts and scoring.',
    'Skill Assessment Module — OpenAI-powered assessments across five tracks with personalized feedback and roadmap recommendations.',
    'Coding Platform Module — Monaco Editor integration with Judge0 compiler support for multi-language code execution.',
    'Contest Management Module — Real-time coding competitions with Socket.IO leaderboards, Elo ratings, and participant management.',
    'Resume Builder Module — Step-by-step resume creation with ATS scanning, multiple templates, and PDF export via html2pdf.js.',
    'Cover Letter Module — Cover letter generation with jsPDF export and localStorage persistence.',
    'Tech Feed Module — News aggregation from NewsAPI with tech news, jobs, internships, and event categories.',
    'Community Forum Module — Q&A forum with question posting, answer threads, and admin moderation capabilities.',
    'Project Recommender Module — Groq AI-powered project suggestions based on user preferences and skill levels.',
    'Notes Module — Comprehensive notes system with AI-generated notes, structured subject content, and user-created notes.',
    'Video Course Module — YouTube API-integrated course platform with progress tracking, bookmarking, and continue-learning features.',
    'Placement Prep Module — Job listings, interview preparation resources, and links to DSA practice areas.',
    'Interview Experience Module — Mock interview interface and interview experience sharing with company and position tagging.',
    'Career Roadmap Module — Visual skill-based and role-based career roadmaps from structured JSON data.',
    'Calendar and Tasks Module — Calendar views with event management and task tracking with Kanban-style views.',
    'Calculator Module — Multi-mode calculator with basic, scientific, graphing, financial, unit conversion, and physics modes.',
    'Premium and Payment Module — Subscription plans with feature comparison and Stripe payment integration.',
    'Admin Panel Module — Comprehensive admin dashboard for user, content, quiz, and analytics management.',
    'Notification Module — In-app notification system for milestones, updates, and platform alerts.',
    'ATS Scanner Module — Resume text analysis for keyword optimization and ATS compatibility scoring.',
    'Global Search Module — Cross-platform search across courses, features, topics, and tools with type-based filters.',
]
for m in modules_list:
    add_bullet(m)

add_page_break()

# ══════════════════════════════════════════════════
# 6. ADVANTAGES / APPLICATIONS
# ══════════════════════════════════════════════════
add_heading_styled('6. Advantages / Applications', 1)

add_subheading('6.1 Advantages')
add_body(
    'The ElevoraAI platform offers a comprehensive set of advantages over existing educational and career development tools:\n'
)
add_bullet('Combines learning, coding practice, resume building, job preparation, and community interaction in a single cohesive platform, eliminating the need for multiple disjointed tools.', 'All-in-One Integration: ')
add_bullet('Integrates Google Gemini, OpenAI GPT-4, and Groq LLaMA 3 to provide personalized study assistance, automated content generation, and intelligent recommendations.', 'Multi-Model AI Integration: ')
add_bullet('Separate interfaces and access controls for users and administrators ensure secure, efficient, and role-appropriate platform usage.', 'Role-Based Architecture: ')
add_bullet('Socket.IO-powered contest rooms with dynamic leaderboards and Elo rating systems provide an engaging competitive learning environment.', 'Real-Time Collaboration: ')
add_bullet('Predefined curricula with module-level progress tracking, streak monitoring, and achievement badges enable systematic skill development.', 'Structured Learning Paths: ')
add_bullet('Built-in ATS scanner, keyword optimization, and PDF export improve resume quality and job application success rates.', 'ATS-Optimized Resume Tools: ')
add_bullet('AI-generated assessments across five tracks with detailed feedback, improvement areas, and personalized roadmap recommendations.', 'Comprehensive Skill Assessment: ')
add_bullet('The AI Study Buddy provides round-the-clock doubt resolution, resource suggestions, and career guidance.', '24x7 AI Assistance: ')
add_bullet('Open-source development model with community contribution guidelines, code of conduct, and transparent development process.', 'Open Source and Community Driven: ')
add_bullet('Admin dashboard provides real-time analytics, user management, content moderation, and system configuration from a single interface.', 'Unified Admin Control: ')
add_bullet('Supports multiple learning styles through video courses, text notes, interactive quizzes, coding practice, and community discussions.', 'Multi-Modal Learning Support: ')
add_bullet('Dark mode, responsive design, toast notifications, and smooth animations provide a modern and accessible user experience.', 'Modern User Experience: ')

add_subheading('6.2 Applications')
add_body(
    'The platform is designed to serve a diverse range of users and use cases:\n'
)
add_bullet('Engineering and computer science students preparing for campus placement drives and technical interviews.', 'College Students: ')
add_bullet('Self-taught programmers and career switchers seeking structured learning paths and practical coding experience.', 'Self-Learners: ')
add_bullet('Working professionals looking to upskill in AI/ML, full-stack development, data science, or cloud technologies.', 'Working Professionals: ')
add_bullet('Colleges and training institutes requiring an integrated platform for course delivery, progress tracking, and student assessment.', 'Educational Institutions: ')
add_bullet('Developers looking to contribute to a real-world open-source project and gain experience in collaborative software development.', 'Open Source Contributors: ')
add_bullet('Coding enthusiasts who want to participate in live contests, compare performance on leaderboards, and improve competitive programming skills.', 'Competitive Programmers: ')
add_bullet('Job seekers who need ATS-optimized resumes, cover letters, interview preparation, and access to curated job opportunities.', 'Job Seekers: ')
add_bullet('Platform administrators and program coordinators who need to manage users, content, and analytics through a unified dashboard.', 'Administrators: ')
add_bullet('Study groups and peer learning communities that benefit from shared notes, forum discussions, and collaborative features.', 'Study Groups: ')
add_bullet('Hackathon organizers who need a platform for contest creation, participant management, and real-time result tracking.', 'Hackathon Organizers: ')

add_page_break()

# ══════════════════════════════════════════════════
# 7. BENEFITS
# ══════════════════════════════════════════════════
add_heading_styled('7. Benefits', 1)

add_body(
    'The ElevoraAI platform delivers a wide range of tangible benefits to its users, administrators, and the broader educational '
    'technology ecosystem. These benefits extend across multiple dimensions including learning efficiency, career readiness, '
    'technological advancement, and community development.'
)

add_body(
    'From a learning efficiency perspective, ElevoraAI eliminates the significant overhead associated with managing multiple '
    'platform accounts, remembering different login credentials, and maintaining separate progress records. When a student uses '
    'Coursera for courses, LeetCode for coding practice, Canva for resume building, and LinkedIn for job searching, they must '
    'constantly switch contexts, adapt to different user interfaces, and manually track their overall progress. ElevoraAI '
    'consolidates all these activities into a single, cohesive experience. The platform\'s unified dashboard provides an '
    'instant overview of all learning activities, streak data, achievements, and pending tasks. This consolidation saves '
    'considerable time and mental energy that can be redirected toward actual learning and skill development. The structured '
    'learning paths with progress tracking ensure that users follow a systematic approach to skill acquisition rather than '
    'jumping between unrelated topics. The streak-based motivation system encourages consistent daily practice, which research '
    'has shown to be far more effective for long-term knowledge retention than sporadic intensive study sessions.'
)

add_body(
    'In terms of career readiness, ElevoraAI provides an end-to-end preparation pipeline that covers the complete journey from '
    'skill acquisition to job placement. The AI-powered skill assessments help users identify their strengths and weaknesses '
    'across multiple domains, allowing them to focus their efforts on areas that need improvement. The personalized learning '
    'roadmaps generated from assessment results provide a clear, actionable path for skill development. The ATS-optimized resume '
    'builder ensures that users\' resumes are formatted and keyword-optimized to pass through automated screening systems, '
    'significantly increasing the likelihood of reaching human recruiters. The cover letter generator complements the resume '
    'with professionally formatted application documents. The placement preparation module provides access to curated job '
    'listings, interview preparation resources, and career roadmaps. The interview experience sharing feature allows users '
    'to learn from the experiences of others who have successfully navigated the interview process at specific companies. '
    'Together, these features create a comprehensive career advancement ecosystem that addresses every stage of the '
    'job-seeking process.'
)

add_body(
    'From a technological standpoint, the platform leverages cutting-edge AI capabilities to enhance learning outcomes. '
    'The multi-model AI architecture — using Google Gemini for conversational assistance, OpenAI for assessment generation '
    'and evaluation, and Groq LLaMA 3 for project recommendations — provides specialized intelligence for different use cases. '
    'This approach ensures that each AI task is handled by the most suitable model, optimizing both response quality and '
    'cost efficiency. The AI Study Buddy is available 24 hours a day, 7 days a week, providing instant assistance whenever '
    'users encounter difficulties. This is particularly valuable for students studying at unconventional hours or those who '
    'may feel hesitant to ask questions in traditional classroom settings. The AI quiz generator can create custom practice '
    'materials on any topic, enabling users to test their knowledge before exams or interviews. The AI note generator produces '
    'well-structured study notes that summarize complex topics into digestible formats.'
)

add_body(
    'For administrators and educational institutions, ElevoraAI provides powerful tools for managing large user bases. '
    'The admin panel offers real-time analytics on platform usage, user engagement, quiz performance, and learning outcomes. '
    'This data-driven approach enables informed decision-making about curriculum design, content strategy, and resource allocation. '
    'The content management system allows administrators to easily add, update, and organize learning materials without requiring '
    'technical expertise. The community moderation tools help maintain a positive and productive learning environment. The system '
    'logs provide a complete audit trail of administrative actions for accountability and security purposes.'
)

add_body(
    'The community and open-source aspects of ElevoraAI deliver additional benefits. The Q&A forum facilitates peer-to-peer '
    'learning, allowing users to benefit from the collective knowledge of the community. The open-source development model '
    'ensures transparency, encourages contributions from developers worldwide, and allows the platform to evolve based on '
    'community needs. Contributors gain real-world experience in collaborative software development, version control, code '
    'review, and project management — skills that are highly valued by employers.'
)

add_body('The key benefits of the platform are listed below:\n')

add_bullet('Reduces the need for multiple competing platforms by providing a fully integrated learning and career advancement solution.')
add_bullet('Saves significant time through AI-powered automation of quiz creation, note generation, and project recommendations.')
add_bullet('Improves learning outcomes through structured curricula, streak-based motivation, achievement badges, and detailed progress tracking.')
add_bullet('Enhances career readiness with ATS-optimized resumes, mock interview tools, skill assessments, and placement preparation resources.')
add_bullet('Provides round-the-clock access to an AI study assistant for instant doubt resolution and personalized guidance.')
add_bullet('Enables administrators to efficiently manage users, content, analytics, and platform settings from a single dashboard.')
add_bullet('Facilitates peer learning and knowledge exchange through the community forum, notes sharing, and collaborative discussions.')
add_bullet('Offers real-time coding competition experience with live leaderboards, Elo rating systems, and multi-language compiler support.')
add_bullet('Supports secure, flexible authentication through email OTP verification, Google OAuth, and JWT-based session management.')
add_bullet('Promotes open-source contribution and collaborative skill development through community-driven development practices.')
add_bullet('Delivers a modern, responsive user experience with dark mode support, smooth animations, and mobile-friendly design.')
add_bullet('Provides multi-modal learning through video courses, text notes, interactive quizzes, coding practice, and community discussions.')

add_page_break()

# ══════════════════════════════════════════════════
# 8. REFERENCES
# ══════════════════════════════════════════════════
add_heading_styled('8. References', 1)

references = [
    'React Documentation. "React 18 — A JavaScript Library for Building User Interfaces." https://react.dev/',
    'Express.js Documentation. "Express — Fast, Unopinionated, Minimalist Web Framework for Node.js." https://expressjs.com/',
    'MongoDB Documentation. "MongoDB Atlas — Global Cloud Database Service." https://www.mongodb.com/atlas',
    'Mongoose Documentation. "Mongoose — Elegant MongoDB Object Modeling for Node.js." https://mongoosejs.com/',
    'OpenAI API Documentation. "GPT-4 API for Natural Language Processing and Code Generation." https://platform.openai.com/',
    'Google Gemini API Documentation. "Gemini — Google\'s Largest and Most Capable AI Model." https://ai.google.dev/',
    'Groq Documentation. "Groq — Fast Inference for LLaMA and Open-Source Language Models." https://console.groq.com/',
    'Socket.IO Documentation. "Socket.IO — Bidirectional and Low-Latency Communication for Every Platform." https://socket.io/',
    'Judge0 API Documentation. "Judge0 — Online Code Execution and Evaluation System." https://judge0.com/',
    'Stripe Documentation. "Stripe — Online Payment Processing for Internet Businesses." https://stripe.com/docs',
    'Tailwind CSS Documentation. "Tailwind CSS — A Utility-First CSS Framework." https://tailwindcss.com/',
    'Vite Documentation. "Vite — Next Generation Frontend Tooling and Build System." https://vitejs.dev/',
    'JWT.io. "JSON Web Tokens — An Open Industry Standard for Secure Data Transmission (RFC 7519)." https://jwt.io/',
    'Monaco Editor Documentation. "Monaco Editor — The Code Editor that Powers Visual Studio Code." https://microsoft.github.io/monaco-editor/',
    'Hamari, J., Koivisto, J., and Sarsa, H. (2014). "Does Gamification Work? — A Literature Review of Empirical Studies on Gamification." Proceedings of the 47th Hawaii International Conference on System Sciences.',
    'Sandhu, R. S., Coyne, E. J., Feinstein, H. L., and Youman, C. E. (1996). "Role-Based Access Control Models." IEEE Computer, 29(2), 38-47.',
    'Pinto, G., Steinmacher, I., and Gerosa, M. A. (2016). "More Common Than You Think: An In-depth Study of Casual Contributors." Proceedings of the IEEE 23rd International Conference on Software Analysis, Evolution, and Reengineering.',
    'Mollick, E. R., and Mollick, L. (2023). "Using AI to Implement Effective Teaching Strategies in Classrooms: Five Strategies, Including Prompts." SSRN Electronic Journal.',
    'Axios Documentation. "Axios — Promise-Based HTTP Client for JavaScript." https://axios-http.com/',
    'Framer Motion Documentation. "Framer Motion — A Production-Ready Animation Library for React." https://www.framer.com/motion/',
]

for i, ref in enumerate(references, 1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.5)
    run = p.add_run(f'[{i}] {ref}')
    run.font.size = Pt(11)

# ── Save .docx ──
docx_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'SYNOPSIS.docx')
doc.save(docx_path)
print(f'DOCX saved: {docx_path}')
